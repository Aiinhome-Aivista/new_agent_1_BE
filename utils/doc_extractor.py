import os
import traceback

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file using pypdf, with error handling."""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        print("Warning: pypdf library is not installed. Unable to parse PDF file natively.")
        return f"[PDF Parsing Fallback: pypdf library is missing. Filename: {os.path.basename(file_path)}]"
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
        traceback.print_exc()
        return f"[Error reading PDF file: {str(e)}]"

def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file using python-docx paragraphs and tables, with error handling."""
    try:
        import docx
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts)
    except ImportError:
        print("Warning: python-docx library is not installed. Unable to parse DOCX file natively.")
        return f"[DOCX Parsing Fallback: python-docx library is missing. Filename: {os.path.basename(file_path)}]"
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
        traceback.print_exc()
        return f"[Error reading DOCX file: {str(e)}]"

def extract_text(file_path):
    """Detects file extension and routes to the appropriate parser, applying preprocessing."""
    if not file_path or not os.path.exists(file_path):
        return ""
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        raw_text = extract_text_from_docx(file_path)
    else:
        # Default text/markdown/json direct reading
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        except Exception as e:
            raw_text = f"[Error reading file as text: {str(e)}]"
            
    # Text preprocessing: cleanup whitespaces and blank lines
    cleaned_lines = []
    for line in raw_text.splitlines():
        trimmed = line.strip()
        if trimmed:
            cleaned_lines.append(trimmed)
            
    return "\n".join(cleaned_lines)
